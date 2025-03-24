from docx import Document
import requests
import json
import time
import re  # Import regex for cleaning text

# Configuration Replace with your actual values
# Ensure you have the required libraries installed: python-docx
DOC_PATH = "your-document.docx"
API_KEY = "your-api-key"
ENDPOINT = "https://your-endpoint.openai.azure.com/openai/deployments/gpt-4/chat/completions?api-version=your-api-version"

HEADERS = {
    "Content-Type": "application/json",
    "api-key": API_KEY,
}

# Load Word Document
doc = Document(DOC_PATH)

# Extract text from paragraphs (Ignoring tables for now)
text = " ".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])  # Ensures one continuous string

# Split text into chunks (to avoid exceeding token limits)
MAX_CHUNK_SIZE = 1000  # Adjust based on token limit for your model the larger the chunk the more data that can be sent
chunks = [text[i:i+MAX_CHUNK_SIZE] for i in range(0, len(text), MAX_CHUNK_SIZE)]

# Function to clean unwanted formatting from summary
def clean_text(text):
    text = re.sub(r"\*\*", "", text)  # Remove all bold formatting (Markdown-style asterisks)
    text = re.sub(r"-\s*", "", text)  # Remove dashes and extra spaces
    text = re.sub(r"\n+", " ", text)  # Remove unnecessary newlines
    text = re.sub(r"\s{2,}", " ", text)  # Remove extra spaces
    return text.strip()

# Function to send request to Azure OpenAI
def send_request(chunk, part_num, total_parts):
    payload = {
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an AI assistant that specializes in summarizing technical reports, test documents, "
                    "and structured data accurately and concisely. Your summaries should be clear, professional, "
                    "and structured logically. Keep the entire summary as one single paragraph, without extra line breaks."
                )
            },
            {
                "role": "user",
                "content": chunk
            }
        ],
        "max_tokens": 300  # Adjust for better summary results, will yield a more detailed summary
    }

    response = requests.post(ENDPOINT, json=payload, headers=HEADERS)

    if response.status_code == 200:
        return clean_text(response.json().get("choices", [{}])[0].get("message", {}).get("content", "No response"))
    elif response.status_code == 429:
        print("Rate limit exceeded. Waiting before retrying...")
        time.sleep(60)  # Increase/Decrease if needed depending on how many requests allowed per minute
        return send_request(chunk, part_num, total_parts)  # Retry the same chunk
    else:
        print(f"Error processing chunk {part_num}/{total_parts}: {response.text}")
        return None

# Sending document in chunks
final_summary = []
for i, chunk in enumerate(chunks):
    print(f"Sending chunk {i+1}/{len(chunks)}...")
    summary_part = send_request(chunk, i+1, len(chunks))
    if summary_part:
        final_summary.append(summary_part)

# Join summary into a single, clean paragraph
summary_text = " ".join(final_summary)

# Append the final summary to the document
doc.add_paragraph("\n")  # Add spacing before summary

final_summary_para = doc.add_paragraph()  # Create a new paragraph
final_summary_para.add_run("Final Summary: ").bold = True  # Bold heading
final_summary_para.add_run(" " + summary_text)  # Add summary text in one paragraph

# Save document with summary appended
NEW_DOC_PATH = DOC_PATH.replace(".docx", "-with-summary.docx")
doc.save(NEW_DOC_PATH)

print(f"\nFinal summary appended to document: {NEW_DOC_PATH}")
